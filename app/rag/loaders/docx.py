"""DOCX loader: headings, paragraphs, and tables in document order.

Headings emerge as `heading` blocks carrying their level — not as
paragraphs — because the heading path (used later for chunk context and
citations) depends on that distinction. Source refs are a running
paragraph/table count: `¶ N`.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document as OpenDocument
from docx.oxml.ns import qn
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph

from app.rag.blocks import Block, LoaderError


class DocxLoader:
    """Loads a `.docx` file into an ordered list of `Block`s."""

    def load(self, path: Path) -> list[Block]:
        path = Path(path)

        try:
            document = OpenDocument(str(path))
        except Exception as exc:
            raise LoaderError(f"could not parse {path.name}: {exc}") from exc

        blocks: list[Block] = []
        ref_counter = 0

        for item in _iter_block_items(document):
            if isinstance(item, DocxParagraph):
                text = item.text.strip()
                if not text:
                    continue
                ref_counter += 1
                ref = f"¶ {ref_counter}"
                level = _heading_level(item.style.name if item.style else None)
                if level is not None:
                    blocks.append(
                        Block(kind="heading", text=text, source_ref=ref, heading_level=level)
                    )
                else:
                    blocks.append(Block(kind="paragraph", text=text, source_ref=ref))
            elif isinstance(item, DocxTable):
                ref_counter += 1
                ref = f"¶ {ref_counter}"
                blocks.append(Block.table(rows=_table_rows(item), source_ref=ref))

        if not blocks:
            raise LoaderError(f"{path.name} contains no extractable content")

        return blocks


def _iter_block_items(document: OpenDocument):
    """Yield paragraphs and tables in the order they appear in the body.

    `document.paragraphs` and `document.tables` are separate flat lists
    that lose their relative order — this walks the underlying XML body
    directly to recover it.
    """
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield DocxParagraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield DocxTable(child, document)


def _heading_level(style_name: str | None) -> int | None:
    if not style_name:
        return None
    if style_name == "Title":
        return 1
    if style_name.startswith("Heading "):
        suffix = style_name.removeprefix("Heading ").strip()
        if suffix.isdigit():
            return int(suffix)
    return None


def _table_rows(table: DocxTable) -> list[list[str]]:
    """The table's cells as a raw grid.

    `cell.text` joins a multi-paragraph cell's paragraphs with newlines,
    so these strings routinely contain "\\n" — `Block.table` normalizes
    that away when it renders, and the structural grid is what every
    downstream row-level decision reads.
    """
    return [[cell.text for cell in row.cells] for row in table.rows]
