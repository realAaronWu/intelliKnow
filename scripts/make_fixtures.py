"""Deterministic synthetic document fixtures for the RAG write-path tests.

Generates ten small documents with exactly known content, used by the
loader tests (increment 03) instead of hand-crafted or downloaded files.
Byte reproducibility is a hard requirement: `duplicate.pdf` must match
`salary_bands.pdf` byte-for-byte, and every fixture must be identical
across repeated runs, so that hash-based tests (this module's own, and
Task 9's duplicate-rejection test) are never flaky.

Run directly to (re)generate the committed fixtures under
`tests/fixtures/docs/`:

    python -m scripts.make_fixtures
"""

from __future__ import annotations

import io
import re
import zipfile
from datetime import datetime
from pathlib import Path

import openpyxl
from docx import Document
from PIL import Image
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen import canvas
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

# ---------------------------------------------------------------------------
# Known values. Tests and loaders import these rather than duplicating them.
# ---------------------------------------------------------------------------

ANNUAL_LEAVE_DAYS = 25

# (band label, min, mid, max) — order matters: rendered top-to-bottom.
SALARY_BANDS: list[tuple[str, int, int, int]] = [
    ("Band 1", 32000, 38000, 44000),
    ("Band 2", 45000, 52000, 59000),
    ("Band 3", 60000, 68000, 76000),
    ("Band 4", 78000, 88000, 98000),
]

REIMBURSEMENT_LIMIT_USD = 150
REIMBURSEMENT_FORM = "Form FIN-204"

# A perfectly rectangular table whose middle column holds *multi-paragraph*
# cells — the single most common real-world table shape that neither
# pdfplumber nor python-docx flattens for you: `cell.text` comes back with
# an embedded "\n". The grid itself is clean (three columns on every row,
# no empty cells), so nothing here should ever be sent to the LLM for
# repair, and no chunk boundary should ever fall inside one of these rows.
# Both properties were broken while table structure was carried only as
# rendered markdown that later code split on "\n".
WRAPPED_TABLE_HEADER = ["Policy", "Detail", "Owner"]

#: (first cell, the detail cell's paragraphs, last cell)
WRAPPED_TABLE_ROWS: list[tuple[str, list[str], str]] = [
    (
        "Annual leave",
        [
            f"Full-time employees accrue {ANNUAL_LEAVE_DAYS} days per calendar year.",
            "Prorated for partial years of service.",
        ],
        "HR",
    ),
    (
        "Expenses",
        [
            f"Meals are reimbursed up to {REIMBURSEMENT_LIMIT_USD} USD per day.",
            f"Submit receipts using {REIMBURSEMENT_FORM}.",
        ],
        "Finance",
    ),
    (
        "Equipment",
        [
            "One laptop per employee, refreshed every three years.",
            "Damaged hardware is replaced within five working days.",
        ],
        "Operations",
    ),
    (
        "Access",
        [
            "VPN access requires manager approval.",
            "Requests are processed within one business day.",
        ],
        "Operations",
    ),
]

BUDGET_SHEET_NAMES = ["Summary", "Q1 Actuals"]
# Cell -> computed value for the formula cells on the "Q1 Actuals" sheet.
# openpyxl never evaluates formulas, so a workbook it writes and that has
# never been opened in Excel has no cached value for them (`data_only=True`
# reads back `None`). These are injected directly into the sheet XML after
# save so the fixture behaves like a workbook Excel has actually computed.
_BUDGET_FORMULA_CACHE = {"C2": 4000, "C3": 4200, "B4": 8200, "C4": 8200}

_FIXED_DATE = datetime(2024, 1, 1)
_FIXED_ZIP_DT = (2024, 1, 1, 0, 0, 0)
_FIXED_MODIFIED_ISO = b"2024-01-01T00:00:00Z"


# ---------------------------------------------------------------------------
# Reproducibility helpers
# ---------------------------------------------------------------------------


def _freeze_zip(path: Path) -> None:
    """Rewrite every entry of a zip-based Office file (.docx/.xlsx) with a
    fixed timestamp, and pin the `dcterms:modified` core property that
    openpyxl/python-docx stamp with the current time on every save.

    Without this, two builds of the same fixture differ only in embedded
    timestamps — byte reproducibility (a hard requirement here) would break.
    """
    data_in = path.read_bytes()
    zin = zipfile.ZipFile(io.BytesIO(data_in), "r")
    out = io.BytesIO()
    zout = zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED)
    for item in zin.infolist():
        data = zin.read(item.filename)
        if item.filename == "docProps/core.xml":
            data = re.sub(
                rb"(<dcterms:modified[^>]*>)[^<]*(</dcterms:modified>)",
                rb"\g<1>" + _FIXED_MODIFIED_ISO + rb"\g<2>",
                data,
            )
        zi = zipfile.ZipInfo(item.filename, date_time=_FIXED_ZIP_DT)
        zi.compress_type = zipfile.ZIP_DEFLATED
        zi.external_attr = item.external_attr
        zout.writestr(zi, data)
    zout.close()
    zin.close()
    path.write_bytes(out.getvalue())


def _inject_formula_cache(path: Path, sheet_arcname: str, values: dict[str, int]) -> None:
    """Patch cached `<v>` values into formula cells of one worksheet part.

    openpyxl writes formula cells as `<c r="C2"><f>B2</f><v></v></c>` with
    an empty cached value. This fills the placeholder in so `data_only=True`
    reads return the computed number rather than `None`.
    """
    data_in = path.read_bytes()
    zin = zipfile.ZipFile(io.BytesIO(data_in), "r")
    out = io.BytesIO()
    zout = zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED)
    refs = "|".join(re.escape(ref) for ref in values)
    pattern = re.compile(
        rb'(<c r="(' + refs.encode() + rb')"[^>]*><f>[^<]*</f><v>)(</v></c>)'
    )

    def repl(match: re.Match) -> bytes:
        ref = match.group(2).decode()
        return match.group(1) + str(values[ref]).encode() + match.group(3)

    for item in zin.infolist():
        data = zin.read(item.filename)
        if item.filename == sheet_arcname:
            data = pattern.sub(repl, data)
        zout.writestr(item, data)
    zout.close()
    zin.close()
    path.write_bytes(out.getvalue())


def _pin_docx_core_properties(document: Document) -> None:
    cp = document.core_properties
    cp.author = "IntelliKnow"
    cp.created = _FIXED_DATE
    cp.modified = _FIXED_DATE
    cp.last_modified_by = "IntelliKnow"
    cp.revision = 1


# ---------------------------------------------------------------------------
# PDF fixtures
# ---------------------------------------------------------------------------


def _salary_table() -> Table:
    header = ["Band", "Min", "Mid", "Max"]
    rows = [header] + [
        [band, str(min_val), str(mid_val), str(max_val)]
        for band, min_val, mid_val, max_val in SALARY_BANDS
    ]
    return Table(
        rows,
        style=TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
            ]
        ),
    )


def _make_handbook(path: Path) -> None:
    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(
        str(path),
        pagesize=letter,
        invariant=1,
        title="Employee Handbook",
        author="IntelliKnow",
    )
    story = [
        Paragraph("Employee Handbook", styles["Title"]),
        Spacer(1, 12),
        Paragraph("Leave Policy", styles["Heading1"]),
        Paragraph(
            f"Full-time employees accrue {ANNUAL_LEAVE_DAYS} days of annual "
            "leave per calendar year, prorated for partial years of service.",
            styles["BodyText"],
        ),
        Spacer(1, 12),
        Paragraph("Compensation", styles["Heading1"]),
        Paragraph(
            "Salary bands are reviewed annually. See the table below.",
            styles["BodyText"],
        ),
        Spacer(1, 12),
        _salary_table(),
    ]
    doc.build(story)


def _make_salary_bands(path: Path) -> None:
    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(
        str(path),
        pagesize=letter,
        invariant=1,
        title="Salary Bands",
        author="IntelliKnow",
    )
    doc.build([Paragraph("Salary Bands", styles["Title"]), Spacer(1, 12), _salary_table()])


def _make_ragged_salary_grid(path: Path) -> None:
    """A merged-cell-style grid: the number of visible columns changes from
    row to row (row 2 merges Band/Min into one wide cell; row 3 merges
    Mid/Max), unlike the uniform grid in `salary_bands.pdf`.
    """
    c = canvas.Canvas(str(path), pagesize=letter, invariant=1)
    c.setTitle("Ragged Salary Grid")
    c.setAuthor("IntelliKnow")
    c.setFont("Helvetica", 10)

    top = 700
    row_h = 20
    left = 50

    def row_y(index: int) -> float:
        return top - index * row_h - 14

    c.drawString(left + 5, row_y(0), "Band")
    c.drawString(left + 105, row_y(0), "Min")
    c.drawString(left + 205, row_y(0), "Mid")
    c.drawString(left + 305, row_y(0), "Max")

    band1 = SALARY_BANDS[0]
    c.drawString(left + 5, row_y(1), f"{band1[0]} (merged)")
    c.drawString(left + 205, row_y(1), str(band1[2]))
    c.drawString(left + 305, row_y(1), str(band1[3]))

    band2 = SALARY_BANDS[1]
    c.drawString(left + 5, row_y(2), band2[0])
    c.drawString(left + 105, row_y(2), str(band2[1]))
    c.drawString(left + 205, row_y(2), f"{band2[2]}-{band2[3]}")

    band3 = SALARY_BANDS[2]
    c.drawString(left + 5, row_y(3), band3[0])
    c.drawString(left + 105, row_y(3), str(band3[1]))
    c.drawString(left + 205, row_y(3), str(band3[2]))
    c.drawString(left + 305, row_y(3), str(band3[3]))

    c.showPage()
    c.save()


def _make_scanned(path: Path) -> None:
    """A valid, image-only PDF page with zero extractable text — deterministic
    pixel data (a formula, not RNG) so the embedded image is reproducible.
    """
    image = Image.new("L", (400, 500), color=200)
    pixels = image.load()
    for x in range(400):
        for y in range(500):
            pixels[x, y] = (x * 3 + y * 7) % 256
    buf = io.BytesIO()
    image.save(buf, format="PNG")
    buf.seek(0)

    c = canvas.Canvas(str(path), pagesize=letter, invariant=1)
    c.setTitle("Scanned Document")
    c.setAuthor("IntelliKnow")
    c.drawImage(ImageReader(buf), 100, 200, width=400, height=500)
    c.showPage()
    c.save()


def _make_corrupt(reference_bytes: bytes, path: Path) -> None:
    """Truncated bytes of a valid, reproducible PDF — deterministic because
    the source is reproducible and the truncation point is fixed.
    """
    path.write_bytes(reference_bytes[: len(reference_bytes) // 2])


# ---------------------------------------------------------------------------
# DOCX fixtures
# ---------------------------------------------------------------------------


def _make_nda(path: Path) -> None:
    document = Document()
    _pin_docx_core_properties(document)

    document.add_heading("Non-Disclosure Agreement", level=1)
    document.add_heading("1. Definitions", level=2)
    document.add_paragraph(
        '"Confidential Information" means any non-public information '
        "disclosed by either party under this Agreement."
    )
    document.add_heading("2. Obligations", level=2)
    document.add_paragraph(
        "1. The Receiving Party shall keep all Confidential Information secret."
    )
    document.add_paragraph(
        "2. The Receiving Party shall not disclose Confidential Information "
        "to any third party without prior written consent."
    )
    document.add_heading("3. Term", level=2)
    table = document.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Effective Date"
    table.rows[0].cells[1].text = "2024-01-01"
    table.rows[1].cells[0].text = "Term (years)"
    table.rows[1].cells[1].text = "2"

    document.save(str(path))
    _freeze_zip(path)


def _make_expense_policy(path: Path) -> None:
    document = Document()
    _pin_docx_core_properties(document)

    document.add_heading("Expense Policy", level=1)
    document.add_paragraph(
        f"Meals during business travel are reimbursed up to "
        f"{REIMBURSEMENT_LIMIT_USD} USD per day. Submit receipts using "
        f"{REIMBURSEMENT_FORM}."
    )
    document.add_heading("Approval", level=2)
    document.add_paragraph(
        "Expenses above the daily limit require manager pre-approval before "
        "the trip begins."
    )

    document.save(str(path))
    _freeze_zip(path)


def _make_wrapped_table(path: Path) -> None:
    """A clean three-column table whose middle cells each hold two
    paragraphs, so `cell.text` reads back with an embedded newline. See
    `WRAPPED_TABLE_ROWS`.
    """
    document = Document()
    _pin_docx_core_properties(document)

    document.add_heading("Policy Summary", level=1)
    table = document.add_table(rows=1 + len(WRAPPED_TABLE_ROWS), cols=3)
    table.style = "Table Grid"

    for column, heading in enumerate(WRAPPED_TABLE_HEADER):
        table.rows[0].cells[column].text = heading

    for index, (policy, detail_paragraphs, owner) in enumerate(WRAPPED_TABLE_ROWS, start=1):
        cells = table.rows[index].cells
        cells[0].text = policy
        first, *rest = detail_paragraphs
        cells[1].text = first
        for paragraph in rest:
            cells[1].add_paragraph(paragraph)
        cells[2].text = owner

    document.save(str(path))
    _freeze_zip(path)


# ---------------------------------------------------------------------------
# XLSX fixture
# ---------------------------------------------------------------------------


def _make_budget(path: Path) -> None:
    workbook = openpyxl.Workbook()
    workbook.properties.creator = "IntelliKnow"
    workbook.properties.created = _FIXED_DATE

    summary = workbook.active
    summary.title = BUDGET_SHEET_NAMES[0]
    summary.append(["Category", "Amount"])
    summary.append(["Travel", 12000])
    summary.append(["Software", 8000])

    q1 = workbook.create_sheet(BUDGET_SHEET_NAMES[1])
    q1.append(["Month", "Spend", "FormulaTotal"])
    q1.append(["Jan", 4000, "=B2"])
    q1.append(["Feb", 4200, "=B3"])
    q1.append(["Total", "=SUM(B2:B3)", "=SUM(C2:C3)"])

    workbook.save(str(path))
    # "Q1 Actuals" is the second sheet created, so it is xl/worksheets/sheet2.xml.
    _inject_formula_cache(path, "xl/worksheets/sheet2.xml", _BUDGET_FORMULA_CACHE)
    _freeze_zip(path)


# ---------------------------------------------------------------------------
# Orchestration
# ---------------------------------------------------------------------------


def build_all(out_dir: Path) -> list[Path]:
    """Generate all ten fixtures into `out_dir` (created if needed).

    Returns the list of written paths. Deterministic: calling this twice
    into different directories produces byte-identical files.
    """
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    handbook_path = out_dir / "handbook.pdf"
    salary_bands_path = out_dir / "salary_bands.pdf"
    ragged_path = out_dir / "ragged_salary_grid.pdf"
    nda_path = out_dir / "nda.docx"
    expense_policy_path = out_dir / "expense_policy.docx"
    wrapped_table_path = out_dir / "wrapped_table.docx"
    budget_path = out_dir / "budget.xlsx"
    corrupt_path = out_dir / "corrupt.pdf"
    scanned_path = out_dir / "scanned.pdf"
    duplicate_path = out_dir / "duplicate.pdf"

    _make_handbook(handbook_path)
    _make_salary_bands(salary_bands_path)
    _make_ragged_salary_grid(ragged_path)
    _make_nda(nda_path)
    _make_expense_policy(expense_policy_path)
    _make_wrapped_table(wrapped_table_path)
    _make_budget(budget_path)
    _make_corrupt(salary_bands_path.read_bytes(), corrupt_path)
    _make_scanned(scanned_path)
    duplicate_path.write_bytes(salary_bands_path.read_bytes())

    return [
        handbook_path,
        salary_bands_path,
        ragged_path,
        nda_path,
        expense_policy_path,
        wrapped_table_path,
        budget_path,
        corrupt_path,
        scanned_path,
        duplicate_path,
    ]


if __name__ == "__main__":
    written = build_all(Path(__file__).resolve().parent.parent / "tests" / "fixtures" / "docs")
    for p in written:
        print(p)
